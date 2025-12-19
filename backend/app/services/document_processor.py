from typing import Dict, Optional, Tuple
import os
from app.services.grammar_service import GrammarService
from app.services.reformulation_service import ReformulationService

# Try to import document loaders
try:
    from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
    DOCUMENT_LOADERS_AVAILABLE = True
except ImportError:
    DOCUMENT_LOADERS_AVAILABLE = False
    PyPDFLoader = None
    TextLoader = None
    Docx2txtLoader = None

# Try to import document generators with structure preservation
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    DOCUMENT_GENERATORS_AVAILABLE = True
except ImportError:
    DOCUMENT_GENERATORS_AVAILABLE = False
    DocxDocument = None

# Try to import python-docx for structure preservation
try:
    from docx import Document as DocxDocumentFull
    from docx.oxml.ns import qn
    DOCX_FULL_AVAILABLE = True
except ImportError:
    DOCX_FULL_AVAILABLE = False
    DocxDocumentFull = None

class DocumentProcessor:
    def __init__(self):
        self.grammar_service = GrammarService()
        self.reformulation_service = ReformulationService()
    
    def extract_text_from_document(self, file_path: str, file_type: str) -> str:
        """
        Extract text from a document.
        
        Args:
            file_path: Path to the document
            file_type: Type of file (pdf, txt, docx)
            
        Returns:
            Extracted text
        """
        if not os.path.exists(file_path):
            raise Exception(f"Le fichier n'existe pas: {file_path}")
        
        if not DOCUMENT_LOADERS_AVAILABLE:
            raise Exception("Document loaders not available. Install langchain-community.")
        
        try:
            if file_type == "pdf" and PyPDFLoader:
                print(f"Loading PDF: {file_path}")
                loader = PyPDFLoader(file_path)
                documents = loader.load()
                print(f"Loaded {len(documents)} pages from PDF")
            elif file_type == "txt" and TextLoader:
                print(f"Loading TXT: {file_path}")
                loader = TextLoader(file_path, encoding="utf-8")
                documents = loader.load()
            elif file_type == "docx" and Docx2txtLoader:
                print(f"Loading DOCX: {file_path}")
                loader = Docx2txtLoader(file_path)
                documents = loader.load()
            else:
                raise Exception(f"Type de fichier non supporté: {file_type}")
            
            if not documents:
                raise Exception("Le document est vide ou ne peut pas être lu.")
            
            # Combine all pages/sections
            full_text = "\n\n".join([doc.page_content for doc in documents if doc.page_content.strip()])
            
            if not full_text or len(full_text.strip()) == 0:
                raise Exception("Aucun texte n'a pu être extrait du document.")
            
            return full_text
        except Exception as e:
            error_msg = str(e)
            print(f"Error extracting text: {error_msg}")
            raise Exception(f"Erreur lors de l'extraction du texte: {error_msg}")
    
    def process_document(self, file_path: str, file_type: str, preserve_structure: bool = True) -> Dict:
        """
        Process a document: extract, correct grammar ONLY (no reformulation to preserve structure).
        
        Args:
            file_path: Path to the document
            file_type: Type of file (pdf, txt, docx)
            preserve_structure: If True, only correct grammar without reformulation to maintain structure
            
        Returns:
            Dictionary with original_text, processed_text, corrections, and statistics
        """
        # Extract text
        print(f"Extracting text from {file_type} file: {file_path}")
        original_text = self.extract_text_from_document(file_path, file_type)
        
        if not original_text or len(original_text.strip()) == 0:
            raise Exception("Le document est vide ou ne contient pas de texte extractible.")
        
        print(f"Extracted {len(original_text)} characters")
        
        # Preserve structure by maintaining line breaks and paragraph separators
        # Split by double newlines to preserve paragraph structure
        paragraphs = original_text.split('\n\n')
        processed_paragraphs = []
        all_corrections = []
        
        print(f"Processing {len(paragraphs)} paragraphs (preserving structure: {preserve_structure})")
        
        # Process paragraphs while preserving structure
        for i, paragraph in enumerate(paragraphs):
            # Preserve empty paragraphs (spacing/structure)
            if not paragraph.strip():
                processed_paragraphs.append(paragraph)
                continue
            
            # Preserve single-line paragraphs (likely headings or titles)
            lines = paragraph.split('\n')
            is_likely_heading = len(lines) == 1 and len(paragraph.strip()) < 100
            
            try:
                # Step 1: Grammar correction ONLY (no reformulation to preserve structure)
                grammar_result = self.grammar_service.correct_text(paragraph)
                corrected_text = grammar_result.get('corrected_text', paragraph)
                corrections = grammar_result.get('corrections', [])
                all_corrections.extend(corrections)
                
                # Only do minimal reformulation if explicitly requested and not a heading
                if not preserve_structure and not is_likely_heading and len(corrected_text) < 500:
                    try:
                        reformulation_result = self.reformulation_service.reformulate_text(
                            corrected_text, 
                            style="academic"
                        )
                        final_text = reformulation_result.get('reformulated_text', corrected_text)
                    except Exception as reform_error:
                        print(f"Reformulation error for paragraph {i}: {reform_error}")
                        final_text = corrected_text
                else:
                    # Preserve structure - only grammar correction
                    final_text = corrected_text
                
                processed_paragraphs.append(final_text)
            except Exception as para_error:
                print(f"Error processing paragraph {i}: {para_error}")
                # If processing fails, use original paragraph to preserve structure
                processed_paragraphs.append(paragraph)
        
        # Reconstruct with original structure (double newlines)
        processed_text = '\n\n'.join(processed_paragraphs)
        
        return {
            "original_text": original_text,
            "processed_text": processed_text,
            "corrections": all_corrections,
            "statistics": {
                "original_length": len(original_text),
                "processed_length": len(processed_text),
                "corrections_count": len(all_corrections),
                "paragraphs_processed": len([p for p in paragraphs if p.strip()])
            }
        }
    
    def generate_document(self, processed_text: str, output_path: str, file_type: str, 
                         original_file_path: Optional[str] = None) -> str:
        """
        Generate a new document with the processed text, preserving structure.
        
        Args:
            processed_text: The processed text to write
            output_path: Path where to save the new document
            file_type: Type of file to generate (pdf, txt, docx)
            original_file_path: Optional path to original document for structure reference
            
        Returns:
            Path to the generated document
        """
        if file_type == "txt":
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(processed_text)
            return output_path
        
        elif file_type == "docx":
            if not DOCUMENT_GENERATORS_AVAILABLE or DocxDocument is None:
                raise Exception("python-docx not available for DOCX generation")
            
            doc = DocxDocument()
            
            # Try to preserve structure from original document if available
            if original_file_path and DOCX_FULL_AVAILABLE and os.path.exists(original_file_path):
                try:
                    # Load original document to get structure
                    original_doc = DocxDocumentFull(original_file_path)
                    paragraphs = processed_text.split('\n\n')
                    para_index = 0
                    
                    # Map processed paragraphs to original structure
                    for orig_para in original_doc.paragraphs:
                        if para_index < len(paragraphs) and paragraphs[para_index].strip():
                            # Use processed text but preserve original formatting
                            new_para = doc.add_paragraph(paragraphs[para_index].strip())
                            
                            # Try to preserve heading style
                            if orig_para.style.name.startswith('Heading'):
                                new_para.style = orig_para.style.name
                            
                            # Preserve alignment
                            if orig_para.alignment:
                                new_para.alignment = orig_para.alignment
                            
                            para_index += 1
                        elif not orig_para.text.strip():
                            # Preserve empty paragraphs for spacing
                            doc.add_paragraph()
                    
                    # Add any remaining processed paragraphs
                    while para_index < len(paragraphs):
                        if paragraphs[para_index].strip():
                            doc.add_paragraph(paragraphs[para_index].strip())
                        para_index += 1
                except Exception as e:
                    print(f"Could not preserve DOCX structure: {e}, using simple structure")
                    # Fallback to simple structure
                    paragraphs = processed_text.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
                        else:
                            doc.add_paragraph()
            else:
                # Simple structure preservation
                paragraphs = processed_text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Detect potential headings (short lines, all caps, or starting with numbers)
                        is_heading = (
                            len(para_text.strip()) < 100 and 
                            (para_text.strip().isupper() or 
                             para_text.strip()[0].isdigit() or
                             para_text.strip().startswith('#') or
                             not '.' in para_text.strip())
                        )
                        
                        para = doc.add_paragraph(para_text.strip())
                        if is_heading:
                            # Apply heading style
                            try:
                                para.style = 'Heading 2'
                            except:
                                pass
                    else:
                        doc.add_paragraph()  # Empty paragraph for spacing
            
            doc.save(output_path)
            return output_path
        
        elif file_type == "pdf":
            if not DOCUMENT_GENERATORS_AVAILABLE:
                raise Exception("reportlab not available for PDF generation")
            
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Create PDF
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split into paragraphs
            paragraphs = processed_text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Clean text for PDF (remove markdown if any)
                    clean_text = para_text.strip().replace('**', '').replace('*', '')
                    p = Paragraph(clean_text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 0.2*inch))
            
            doc.build(story)
            return output_path
        
        else:
            raise Exception(f"Unsupported output file type: {file_type}")
    
    def process_and_generate(self, file_path: str, file_type: str, output_dir: str = "./processed_documents") -> Dict:
        """
        Complete pipeline: extract, process, and generate new document.
        
        Args:
            file_path: Path to original document
            file_type: Type of file
            output_dir: Directory to save processed document
            
        Returns:
            Dictionary with all results and path to generated document
        """
        os.makedirs(output_dir, exist_ok=True)
        
        # Process document
        process_result = self.process_document(file_path, file_type)
        
        # Generate output filename
        original_filename = os.path.basename(file_path)
        name_without_ext = os.path.splitext(original_filename)[0]
        output_filename = f"{name_without_ext}_amélioré.{file_type}"
        output_path = os.path.join(output_dir, output_filename)
        
        # Generate new document with structure preservation
        generated_path = self.generate_document(
            process_result['processed_text'],
            output_path,
            file_type,
            original_file_path=file_path  # Pass original file for structure reference
        )
        
        return {
            **process_result,
            "generated_document_path": generated_path,
            "generated_document_filename": output_filename
        }

